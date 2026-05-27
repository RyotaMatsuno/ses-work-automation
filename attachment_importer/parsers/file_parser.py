# -*- coding: utf-8 -*-
"""ファイル（Excel/Word/PDF）からテキストを抽出し、人員情報を構造化する。"""

import os, re
from pathlib import Path


def extract_text_from_excel(filepath: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(filepath, data_only=True)
    texts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            row_texts = [str(c) for c in row if c is not None and str(c).strip()]
            if row_texts:
                texts.append(' '.join(row_texts))
    return '\n'.join(texts)


def extract_text_from_word(filepath: str) -> str:
    from docx import Document
    doc = Document(filepath)
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_texts = [c.text.strip() for c in row.cells if c.text.strip()]
            if row_texts:
                texts.append(' | '.join(row_texts))
    return '\n'.join(texts)


def extract_text_from_pdf(filepath: str) -> str:
    import pdfplumber
    texts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
    return '\n'.join(texts)


def extract_text_from_file(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    if ext in ('.xlsx', '.xls'):
        return extract_text_from_excel(filepath)
    elif ext in ('.docx', '.doc'):
        return extract_text_from_word(filepath)
    elif ext == '.pdf':
        return extract_text_from_pdf(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def parse_file(filepath: str, api_key: str = None) -> list:
    """ファイルを解析して人員情報リストを返す。"""
    from parsers.text_parser import parse_text
    text = extract_text_from_file(filepath)
    if not text.strip():
        return []
    return parse_text(text, api_key)
