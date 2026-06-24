# -*- coding: utf-8 -*-
import base64
import io
import json
import os
import re
from datetime import datetime

from dotenv import dotenv_values

from line_bridge import guarded_anthropic_call

ENV_PATH = os.path.join(os.path.dirname(__file__), "..", "config", ".env")
if os.path.exists(ENV_PATH):
    config = dotenv_values(ENV_PATH)
    for key, value in config.items():
        if key not in os.environ:
            os.environ[key] = value

ANTHROPIC_API_KEY = os.environ.get("ANTHROPIC_API_KEY", "")


def _extract_pdf_text(file_bytes):
    try:
        import pdfplumber

        texts = []
        with pdfplumber.open(io.BytesIO(file_bytes)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    texts.append(text)
        return "\n".join(texts)
    except Exception as e:
        print(f"[skill_extractor] pdf extract error: {e}")
        return ""


def _extract_xlsx_text(file_bytes):
    try:
        from openpyxl import load_workbook

        texts = []
        workbook = load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
        try:
            for sheet in workbook.worksheets:
                texts.append(f"[sheet: {sheet.title}]")
                for row in sheet.iter_rows(values_only=True):
                    values = [str(v) for v in row if v is not None and str(v).strip()]
                    if values:
                        texts.append("\t".join(values))
        finally:
            workbook.close()
        return "\n".join(texts)
    except Exception as e:
        print(f"[skill_extractor] xlsx extract error: {e}")
        return ""


def _extract_docx_text(file_bytes):
    try:
        from docx import Document

        document = Document(io.BytesIO(file_bytes))
        texts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
                if values:
                    texts.append("\t".join(values))
        return "\n".join(texts)
    except Exception as e:
        print(f"[skill_extractor] docx extract error: {e}")
        return ""


def extract_skills_from_bytes(file_bytes, mime_type):
    try:
        mime_type = mime_type or ""
        if mime_type == "application/pdf":
            return _extract_pdf_text(file_bytes)
        if mime_type == "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet":
            return _extract_xlsx_text(file_bytes)
        if mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return _extract_docx_text(file_bytes)
        if mime_type.startswith("image/"):
            return base64.b64encode(file_bytes).decode("utf-8")
        for extractor in (_extract_pdf_text, _extract_xlsx_text, _extract_docx_text):
            text = extractor(file_bytes)
            if text.strip():
                return text
        return file_bytes.decode("utf-8", errors="ignore")
    except Exception as e:
        print(f"[skill_extractor] extract error: {e}")
        return ""


def _parse_json_text(text):
    try:
        cleaned = re.sub(r"```json|```", "", text or "").strip()
        match = re.search(r"\{.*\}", cleaned, re.S)
        if match:
            cleaned = match.group(0)
        data = json.loads(cleaned)
        return data if isinstance(data, dict) else {}
    except Exception as e:
        print(f"[skill_extractor] json parse error: {e}")
        return {}


def analyze_skill_sheet_v2(file_bytes, mime_type, summary_text=""):
    try:
        extracted = extract_skills_from_bytes(file_bytes, mime_type)
        if not extracted:
            return {}

        current_year = datetime.now().year
        cutoff_year = current_year - 10

        system = (
            "SES\u696d\u754c\u306e\u30b9\u30ad\u30eb\u30b7\u30fc\u30c8\u89e3\u6790AI\u3002"
            "JSON\u5f62\u5f0f\u306e\u307f\u8fd4\u7b54\u3002"
            "price=\u4e07\u5186\u6574\u6570\u3002experience_years=\u696d\u754c\u7d4c\u9a13\u5e74\u6570\u306e\u6574\u6570\u3002"
        )

        prompt = f"""SES\u696d\u754c\u306e\u30b9\u30ad\u30eb\u30b7\u30fc\u30c8\u304b\u3089\u4eba\u6750\u60c5\u5831\u3092JSON\u306e\u307f\u3067\u62bd\u51fa\u3057\u3066\u304f\u3060\u3055\u3044\u3002\u30de\u30fc\u30af\u30c0\u30a6\u30f3\u4e0d\u8981\u3002

\u73fe\u5728\u5e74: {current_year}
10\u5e74\u4ee5\u4e0a\u672a\u4f7f\u7528\u5224\u5b9a\u5e74: {cutoff_year}\u5e74\u4ee5\u524d

\u51fa\u529b\u5f62\u5f0f:
{{"name":"","age":0,"experience_years":0,"price":0,"available_date":"YYYY-MM-DD or sokujitsu","location":"","affiliation":"","note":"","skills":[{{"name":"Java","years":20,"last_used_year":2025,"active":true}}]}}

\u30b9\u30ad\u30eb\u306eactive\u5224\u5b9a\u30eb\u30fc\u30eb:
- last_used_year\u304c{cutoff_year}\u5e74\u4ee5\u524d \u2192 active=false
- last_used_year\u4e0d\u660e\u3067\u3001\u8077\u6b74\u306b\u4e00\u5ea6\u3082\u767b\u5834\u3057\u306a\u3044 \u2192 active=false
- years\u304c1\u5e74\u672a\u6e80 \u2192 active=false
- \u4e0a\u8a18\u4ee5\u5916 \u2192 active=true

summary_text\u304c\u3042\u308c\u3070\u3001\u30b9\u30ad\u30eb\u30b7\u30fc\u30c8\u672c\u6587\u3092\u512a\u5148\u3057\u3064\u3064\u60f3\u5b9a\u5358\u4fa1\u30fb\u7a3c\u50cd\u5f00\u59cb\u65e5\u30fb\u5c45\u4f4f\u5730\u306a\u3069\u3092\u88dc\u5b8c\u3057\u3066\u304f\u3060\u3055\u3044\u3002

summary_text:
{summary_text or "(\u306a\u3057)"}
"""

        if mime_type and mime_type.startswith("image/"):
            content = [
                {"type": "image", "source": {"type": "base64", "media_type": mime_type, "data": extracted}},
                {
                    "type": "text",
                    "text": f"{prompt}\n\n\u30b9\u30ad\u30eb\u30b7\u30fc\u30c8\u5185\u5bb9\u306f\u753b\u50cf\u3092\u78ba\u8a8d\u3057\u3066\u304f\u3060\u3055\u3044\u3002",
                },
            ]
        else:
            combined_text = "\n\n".join([t for t in [summary_text, extracted] if t])
            content = f"{prompt}\n\n\u30b9\u30ad\u30eb\u30b7\u30fc\u30c8\u5185\u5bb9:\n{combined_text[:50000]}"

        result = guarded_anthropic_call(
            system,
            content,
            max_tokens=2000,
            caller="analyze_skill_sheet_v2",
            model="claude-haiku-4-5-20251001",
        )
        return _parse_json_text(result)
    except Exception as e:
        print(f"[skill_extractor] analyze error: {e}")
        return {}


def filter_and_sort_skills(skills_list):
    try:
        from webhook_server import VALID_SKILLS

        valid_map = {skill.lower(): skill for skill in VALID_SKILLS}
        active_skills = []
        for item in skills_list or []:
            if not isinstance(item, dict) or not item.get("active"):
                continue
            name = str(item.get("name") or "").strip()
            normalized = valid_map.get(name.lower())
            if not normalized:
                continue
            try:
                years = float(item.get("years") or 0)
            except Exception:
                years = 0
            active_skills.append((normalized, years))

        sorted_names = []
        seen = set()
        for name, _ in sorted(active_skills, key=lambda x: x[1], reverse=True):
            if name not in seen:
                sorted_names.append(name)
                seen.add(name)
        return sorted_names
    except Exception as e:
        print(f"[skill_extractor] filter error: {e}")
        return []
