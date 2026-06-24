# -*- coding: utf-8 -*-
"""
attachment_importer 全ファイル生成スクリプト
ses_work/ から実行: python build_importer.py
"""

import os

BASE = os.path.join(os.path.dirname(__file__), "attachment_importer")

FILES = {}

# ============================================================
# parsers/text_parser.py
# ============================================================
FILES["parsers/text_parser.py"] = r'''# -*- coding: utf-8 -*-
"""テキストサマリーから人員情報を構造化抽出する。"""

import json, os, re, sys
from datetime import date
from typing import Any
import anthropic

DELIMITER_PATTERNS = [r'-{10,}', r'ー{5,}', r'━{5,}', r'={10,}', r'_{10,}', r'\*{10,}']
DELIMITER_RE = re.compile('|'.join(DELIMITER_PATTERNS))
NAME_KEYWORDS = ['氏名', '名前', '■名前', '【名前】', '【氏名】', '■氏名', 'お名前']
INITIAL_RE = re.compile(r'[A-Z]\.[A-Z]')


def split_into_blocks(text: str) -> list:
    parts = DELIMITER_RE.split(text)
    blocks = []
    for part in parts:
        part = part.strip()
        if not part:
            continue
        has_name = any(kw in part for kw in NAME_KEYWORDS)
        has_initial = bool(INITIAL_RE.search(part))
        if has_name or has_initial:
            blocks.append(part)
    return blocks


def extract_person_from_block(block: str, client, today: str) -> dict:
    prompt = f"""以下の人員情報テキストから構造化データを抽出してください。
不明な項目はnullにしてください。JSONのみ返してください（マークダウン不要）。
今日の日付: {today}

{block}

返すJSON:
{{
  "name": "氏名",
  "age": null,
  "gender": null,
  "nearest_station": null,
  "affiliation": null,
  "price": null,
  "available_date": null,
  "skills_list": [],
  "experience_years": null,
  "contact_email": null,
  "contact_phone": null,
  "remote_preference": null,
  "raw_text": "原文全体"
}}"""
    for attempt in range(3):
        try:
            msg = client.messages.create(
                model="claude-haiku-4-5-20251001",
                max_tokens=2000,
                messages=[{"role": "user", "content": prompt}],
            )
            raw = msg.content[0].text.strip()
            raw = re.sub(r'^```[a-z]*\n?', '', raw)
            raw = re.sub(r'\n?```$', '', raw)
            return json.loads(raw)
        except Exception as e:
            if attempt == 2:
                return {"name": None, "raw_text": block, "_parse_error": str(e)}
    return {"name": None, "raw_text": block}


def parse_text(text: str, api_key: str = None) -> list:
    key = api_key or os.environ.get("ANTHROPIC_API_KEY", "")
    client = anthropic.Anthropic(api_key=key)
    today = date.today().isoformat()
    blocks = split_into_blocks(text)
    if not blocks:
        blocks = [text.strip()]
    results = []
    for block in blocks:
        person = extract_person_from_block(block, client, today)
        if person.get("name") or person.get("raw_text"):
            results.append(person)
    return results
'''

# ============================================================
# parsers/file_parser.py
# ============================================================
FILES["parsers/file_parser.py"] = r'''# -*- coding: utf-8 -*-
"""ファイル（Excel/Word/PDF）からテキストを抽出し、人員情報を構造化する。"""

import os, re
from pathlib import Path


def extract_text_from_excel(filepath: str) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(filepath, data_only=True)
    texts = []
    for ws in wb.worksheets:
        for row in ws.iter_rows(values_only=True):
            row_texts = [str(c) for c in row if c is not None and str(c).strip()]
            if row_texts:
                texts.append(' '.join(row_texts))
    return '\n'.join(texts)


def extract_text_from_word(filepath: str) -> str:
    from docx import Document
    doc = Document(filepath)
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_texts = [c.text.strip() for c in row.cells if c.text.strip()]
            if row_texts:
                texts.append(' | '.join(row_texts))
    return '\n'.join(texts)


def extract_text_from_pdf(filepath: str) -> str:
    import pdfplumber
    texts = []
    with pdfplumber.open(filepath) as pdf:
        for page in pdf.pages:
            t = page.extract_text()
            if t:
                texts.append(t)
    return '\n'.join(texts)


def extract_text_from_file(filepath: str) -> str:
    ext = Path(filepath).suffix.lower()
    if ext in ('.xlsx', '.xls'):
        return extract_text_from_excel(filepath)
    elif ext in ('.docx', '.doc'):
        return extract_text_from_word(filepath)
    elif ext == '.pdf':
        return extract_text_from_pdf(filepath)
    else:
        raise ValueError(f"Unsupported file type: {ext}")


def parse_file(filepath: str, api_key: str = None) -> list:
    """ファイルを解析して人員情報リストを返す。"""
    from parsers.text_parser import parse_text
    text = extract_text_from_file(filepath)
    if not text.strip():
        return []
    return parse_text(text, api_key)
'''

# ============================================================
# utils/drive_downloader.py
# ============================================================
FILES["utils/drive_downloader.py"] = r'''# -*- coding: utf-8 -*-
"""Google SpreadsheetをExcelとしてダウンロードする。"""

import os, re, requests
from pathlib import Path


def extract_file_id(url: str) -> str:
    m = re.search(r'/spreadsheets/d/([a-zA-Z0-9_-]+)', url)
    if m:
        return m.group(1)
    raise ValueError(f"SpreadsheetのURLからIDを抽出できませんでした: {url}")


def download_spreadsheet(url: str, save_dir: str) -> str:
    """
    GoogleスプレッドシートをExcel形式でダウンロードし、保存パスを返す。
    認証不要の公開シートのみ対応。認証が必要な場合はValueErrorを送出。
    """
    file_id = extract_file_id(url)
    export_url = f"https://docs.google.com/spreadsheets/d/{file_id}/export?format=xlsx"
    resp = requests.get(export_url, timeout=30, allow_redirects=True)
    if resp.status_code == 401 or 'accounts.google.com' in resp.url:
        raise ValueError("このスプレッドシートは認証が必要です。URLのみNotionに保存します。")
    resp.raise_for_status()
    os.makedirs(save_dir, exist_ok=True)
    save_path = os.path.join(save_dir, f"{file_id}.xlsx")
    with open(save_path, 'wb') as f:
        f.write(resp.content)
    return save_path
'''

# ============================================================
# utils/notion_writer.py
# ============================================================
FILES["utils/notion_writer.py"] = r'''# -*- coding: utf-8 -*-
"""NotionエンジニアDBへの登録・更新処理。"""

import json, os, re, requests
from datetime import date

NOTION_VERSION = "2022-06-28"
ENGINEER_DB_ID = "343450ff-37c0-819d-8769-fb0a8a4ceeb1"

# エンジニアDBに存在するスキルオプション（確認済み）
VALID_SKILLS = [
    "Java", "Python", "PHP", "JavaScript", "TypeScript", "C#", "Node.js",
    "React", "AWS", "インフラ", "PostgreSQL", "Oracle", "Vue.js", "MySQL",
    "Swift", "Azure", "Linux", "Go", "Ruby", "Docker", "MongoDB", "Spring", "SQL Server",
]

# スキル名の正規化マッピング（よくある表記ゆれ）
SKILL_NORMALIZE = {
    "java": "Java", "python": "Python", "php": "PHP",
    "javascript": "JavaScript", "js": "JavaScript",
    "typescript": "TypeScript", "ts": "TypeScript",
    "c#": "C#", "nodejs": "Node.js", "node.js": "Node.js",
    "react": "React", "reactjs": "React",
    "aws": "AWS", "postgresql": "PostgreSQL", "postgre": "PostgreSQL",
    "oracle": "Oracle", "vuejs": "Vue.js", "vue.js": "Vue.js", "vue": "Vue.js",
    "mysql": "MySQL", "swift": "Swift", "azure": "Azure",
    "linux": "Linux", "go": "Go", "ruby": "Ruby",
    "docker": "Docker", "mongodb": "MongoDB",
    "spring": "Spring", "sqlserver": "SQL Server", "sql server": "SQL Server",
    "pl/sql": "Oracle", "plsql": "Oracle",
    "c++": "C#",  # 近似マッピング（C++はDBにないためC#に寄せず備考行き）
}


def normalize_skills(skills_list: list) -> tuple:
    """
    スキルリストをDBオプションに正規化する。
    Returns: (valid_skills, unknown_skills)
    """
    valid, unknown = [], []
    for s in (skills_list or []):
        normalized = SKILL_NORMALIZE.get(s.lower(), s)
        if normalized in VALID_SKILLS:
            valid.append(normalized)
        else:
            unknown.append(s)
    return list(set(valid)), unknown


def source_to_assignee(source: str) -> str:
    if "松野" in source:
        return "松野"
    if "岡本" in source:
        return "岡本"
    return "共通"


def search_engineer_by_name(name: str, headers: dict) -> str | None:
    """名前でエンジニアDBを検索し、既存ページIDを返す。なければNone。"""
    url = f"https://api.notion.com/v1/databases/{ENGINEER_DB_ID}/query"
    payload = {
        "filter": {
            "property": "名前",
            "title": {"equals": name}
        }
    }
    resp = requests.post(url, headers=headers, json=payload, timeout=30)
    resp.raise_for_status()
    results = resp.json().get("results", [])
    return results[0]["id"] if results else None


def build_properties(record: dict, source: str) -> dict:
    """Notionプロパティ辞書を構築する。"""
    valid_skills, unknown_skills = normalize_skills(record.get("skills_list") or [])
    assignee = source_to_assignee(source)

    props = {
        "名前": {"title": [{"text": {"content": record.get("name") or "（名前不明）"}}]},
        "担当者": {"select": {"name": assignee}},
        "入力元": {"select": {"name": source}},
        "稼働状況": {"select": {"name": "稼働可能"}},
    }

    if valid_skills:
        props["スキル"] = {"multi_select": [{"name": s} for s in valid_skills]}

    if record.get("price"):
        try:
            props["単価（万円）"] = {"number": float(record["price"])}
        except (ValueError, TypeError):
            pass

    if record.get("available_date"):
        try:
            d = str(record["available_date"])[:10]
            date.fromisoformat(d)  # バリデーション
            props["稼働可能日"] = {"date": {"start": d}}
        except (ValueError, TypeError):
            pass

    if record.get("nearest_station"):
        props["最寄り駅"] = {"rich_text": [{"text": {"content": str(record["nearest_station"])[:200]}}]}

    if record.get("affiliation"):
        aff = str(record["affiliation"])[:200]
        props["所属会社名"] = {"rich_text": [{"text": {"content": aff}}]}
        props["所属会社"] = {"rich_text": [{"text": {"content": aff}}]}

    if record.get("experience_years"):
        try:
            props["経験年数"] = {"number": float(record["experience_years"])}
        except (ValueError, TypeError):
            pass

    if record.get("contact_email"):
        props["メール"] = {"email": str(record["contact_email"])}

    if record.get("contact_phone"):
        props["連絡先"] = {"phone_number": str(record["contact_phone"])}

    # イニシャル自動生成
    name = record.get("name") or ""
    m = re.search(r'([A-Z])\.([A-Z])', name)
    if m:
        props["イニシャル"] = {"rich_text": [{"text": {"content": m.group(0)}}]}

    # 人員情報原文
    raw = record.get("raw_text") or ""
    if raw:
        props["人員情報原文"] = {"rich_text": [{"text": {"content": raw[:2000]}}]}

    # 不明スキルを備考に追記
    if unknown_skills:
        note = "【スキル追記】" + ", ".join(unknown_skills)
        props["備考（LINEメモ）"] = {"rich_text": [{"text": {"content": note[:2000]}}]}

    return props


def upsert_engineer(record: dict, source: str, file_path: str = None,
                    drive_url: str = None, headers: dict = None, dry_run: bool = False) -> dict:
    """
    エンジニアDBにレコードを登録または更新する。
    Returns: {"action": "create/update/skip", "page_id": str, "name": str}
    """
    if not headers:
        raise ValueError("notion headers required")

    name = record.get("name")
    if not name:
        return {"action": "skip", "reason": "name is empty"}

    props = build_properties(record, source)

    # ファイルパス / DriveURL
    if file_path:
        props["添付ファイルパス"] = {"rich_text": [{"text": {"content": str(file_path)[:2000]}}]}
    if drive_url:
        props["DriveリンクURL"] = {"url": str(drive_url)}

    if dry_run:
        return {"action": "dry_run", "name": name, "properties": props}

    existing_id = search_engineer_by_name(name, headers)

    if existing_id:
        url = f"https://api.notion.com/v1/pages/{existing_id}"
        resp = requests.patch(url, headers=headers, json={"properties": props}, timeout=30)
        resp.raise_for_status()
        return {"action": "update", "page_id": existing_id, "name": name}
    else:
        url = "https://api.notion.com/v1/pages"
        payload = {
            "parent": {"database_id": ENGINEER_DB_ID},
            "properties": props,
        }
        resp = requests.post(url, headers=headers, json=payload, timeout=30)
        resp.raise_for_status()
        page_id = resp.json().get("id", "")
        return {"action": "create", "page_id": page_id, "name": name}
'''

# ============================================================
# importer.py (メイン)
# ============================================================
FILES["importer.py"] = r'''# -*- coding: utf-8 -*-
"""
attachment_importer - メインスクリプト。
人員情報テキスト + ファイルを解析してNotionエンジニアDBに登録する。

使用例:
  python attachment_importer/importer.py --text "..." --file sheet.xlsx --source "松野LINE"
  python attachment_importer/importer.py --text "..." --spreadsheet-url "https://..." --source "岡本メール"
  python attachment_importer/importer.py --text "..." --file sheet.xlsx --source "松野LINE" --dry-run
"""

import argparse, json, logging, os, sys
from datetime import datetime
from pathlib import Path

# パス解決
BASE_DIR = Path(__file__).parent
SES_WORK_DIR = BASE_DIR.parent
sys.path.insert(0, str(SES_WORK_DIR))

try:
    from dotenv import dotenv_values
    _env = dotenv_values(SES_WORK_DIR / "config" / ".env")
    for k, v in _env.items():
        if k not in os.environ and v:
            os.environ[k] = v
except Exception:
    pass

from attachment_importer.parsers.text_parser import parse_text
from attachment_importer.parsers.file_parser import parse_file, extract_text_from_file
from attachment_importer.utils.drive_downloader import download_spreadsheet
from attachment_importer.utils.notion_writer import upsert_engineer

import requests

NOTION_VERSION = "2022-06-28"
LOG_PATH = BASE_DIR / "import.log"
FAILED_PATH = BASE_DIR / "failed_imports.json"
DOWNLOAD_DIR = str(BASE_DIR / "downloaded_files")

logging.basicConfig(
    level=logging.INFO,
    format="[%(asctime)s] %(levelname)s %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
    handlers=[
        logging.FileHandler(LOG_PATH, encoding="utf-8"),
        logging.StreamHandler(sys.stdout),
    ],
)
logger = logging.getLogger(__name__)


def build_notion_headers() -> dict:
    api_key = os.environ.get("NOTION_API_KEY", "")
    if not api_key:
        raise RuntimeError("NOTION_API_KEY が設定されていません")
    return {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json",
        "Notion-Version": NOTION_VERSION,
    }


def merge_records(text_records: list, file_records: list) -> list:
    """テキストとファイルの人員情報を氏名マッチングでマージ。"""
    if not file_records:
        return text_records
    if not text_records:
        return file_records

    merged = []
    used_file_idx = set()

    for tr in text_records:
        tname = (tr.get("name") or "").strip()
        best_idx, best_score = None, 0

        for i, fr in enumerate(file_records):
            if i in used_file_idx:
                continue
            fname = (fr.get("name") or "").strip()
            if not fname:
                continue
            # スコアリング: 完全一致=3, イニシャル一致=2, 部分一致=1
            if tname == fname:
                score = 3
            elif tname and fname and (tname in fname or fname in tname):
                score = 2
            elif tname and fname and tname[:3] == fname[:3]:
                score = 1
            else:
                score = 0

            if score > best_score:
                best_score, best_idx = score, i

        if best_idx is not None and best_score > 0:
            fr = file_records[best_idx]
            used_file_idx.add(best_idx)
            # ファイル情報優先でマージ
            combined = {**tr, **{k: v for k, v in fr.items() if v is not None}}
            combined["raw_text"] = tr.get("raw_text") or fr.get("raw_text") or ""
            merged.append(combined)
        else:
            merged.append(tr)

    # マッチしなかったファイルレコードを追加
    for i, fr in enumerate(file_records):
        if i not in used_file_idx and fr.get("name"):
            merged.append(fr)

    return merged


def run(text: str, file_path: str = None, spreadsheet_url: str = None,
        source: str = "松野LINE", dry_run: bool = False):
    """メイン処理。"""
    api_key = os.environ.get("ANTHROPIC_API_KEY", "")
    notion_headers = build_notion_headers()
    failed = []

    # テキスト解析
    logger.info(f"テキスト解析開始 source={source}")
    text_records = []
    try:
        text_records = parse_text(text, api_key)
        logger.info(f"テキスト解析完了: {len(text_records)}人分")
    except Exception as e:
        logger.error(f"テキスト解析失敗: {e}")

    # ファイル解析
    file_records = []
    actual_file_path = file_path
    drive_url = None

    if spreadsheet_url:
        drive_url = spreadsheet_url
        try:
            actual_file_path = download_spreadsheet(spreadsheet_url, DOWNLOAD_DIR)
            logger.info(f"スプレッドシートDL完了: {actual_file_path}")
        except Exception as e:
            logger.warning(f"スプレッドシートDL失敗（URLのみ保存）: {e}")
            actual_file_path = None

    if actual_file_path and os.path.exists(actual_file_path):
        try:
            file_records = parse_file(actual_file_path, api_key)
            logger.info(f"ファイル解析完了: {len(file_records)}人分 ({actual_file_path})")
        except Exception as e:
            logger.warning(f"ファイル解析失敗、テキスト情報のみで続行: {e}")

    # マージ
    records = merge_records(text_records, file_records)
    logger.info(f"マージ完了: {len(records)}人分")

    # Notion登録
    success_count = 0
    for record in records:
        name = record.get("name") or "（名前不明）"
        try:
            result = upsert_engineer(
                record=record,
                source=source,
                file_path=actual_file_path,
                drive_url=drive_url,
                headers=notion_headers,
                dry_run=dry_run,
            )
            action = result.get("action")
            if action == "dry_run":
                logger.info(f"[DRY-RUN] {name} → 登録スキップ スキル={[s['name'] for s in result.get('properties', {}).get('スキル', {}).get('multi_select', [])]}")
            elif action == "create":
                logger.info(f"[SUCCESS] {name} → 新規登録 page_id={result.get('page_id')}")
            elif action == "update":
                logger.info(f"[SUCCESS] {name} → 更新 page_id={result.get('page_id')}")
            elif action == "skip":
                logger.warning(f"[SKIP] {name} → {result.get('reason')}")
            success_count += 1
        except Exception as e:
            logger.error(f"[ERROR] {name} → Notion登録失敗: {e}")
            failed.append({"name": name, "error": str(e), "record": record})

    # 失敗ログ保存
    if failed:
        existing = []
        if FAILED_PATH.exists():
            try:
                existing = json.loads(FAILED_PATH.read_text(encoding="utf-8"))
            except Exception:
                pass
        existing.extend(failed)
        FAILED_PATH.write_text(json.dumps(existing, ensure_ascii=False, indent=2), encoding="utf-8")
        logger.warning(f"{len(failed)}件の登録に失敗しました → {FAILED_PATH}")

    logger.info(f"完了: 成功={success_count} 失敗={len(failed)} 合計={len(records)}")
    return {"success": success_count, "failed": len(failed), "total": len(records)}


def main():
    parser = argparse.ArgumentParser(description="人員情報をNotionエンジニアDBに登録する")
    parser.add_argument("--text", required=True, help="人員情報テキスト（LINEメッセージ等）")
    parser.add_argument("--file", help="スキルシートファイルパス（.xlsx/.xls/.docx/.pdf）")
    parser.add_argument("--spreadsheet-url", help="GoogleスプレッドシートURL")
    parser.add_argument("--source", default="松野LINE",
                        choices=["松野LINE", "岡本LINE", "松野メール", "岡本メール", "共通メール"],
                        help="入力元ラベル")
    parser.add_argument("--dry-run", action="store_true", help="Notionに書かずログだけ出力")
    args = parser.parse_args()

    run(
        text=args.text,
        file_path=args.file,
        spreadsheet_url=args.spreadsheet_url,
        source=args.source,
        dry_run=args.dry_run,
    )


if __name__ == "__main__":
    main()
'''

# ============================================================
# __init__.py
# ============================================================
FILES["__init__.py"] = ""
FILES["parsers/__init__.py"] = ""
FILES["utils/__init__.py"] = ""
FILES["tests/__init__.py"] = ""

# ============================================================
# tests/test_text_parser.py
# ============================================================
FILES["tests/test_text_parser.py"] = r'''# -*- coding: utf-8 -*-
"""text_parserの単体テスト（区切り線分割のみ、API不要）"""

import sys, os
sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..', '..'))

from attachment_importer.parsers.text_parser import split_into_blocks

SAMPLE_1 = """
おつかれさまです！現在5月から案件でお世話になっています。
------------------------------
【名　前】H.S（55歳/男性）
【最寄駅】北小金駅
【稼　働】7月～
【単　金】70万
【スキル】Java、Spring、JavaScript
------------------------------
"""

SAMPLE_2 = """
@All お世話になっております。3名営業しています。
ーーーーーーーーーーーーー
■氏名：OA（33歳・女性）
■最寄り駅：森林公園駅
■希望単価：60万
■スキル概要：Java,C#,React
ーーーーーーーーーーーーー
ーーーーーーーーーーーーー
氏名：R.H（男性／24歳）
最寄駅：学芸大学駅
希望単価：45万～50万円
スキル：Java／PostgreSQL
ーーーーーーーーーーーーー
ーーーーーーーーーーーーー
氏名:U.H／33歳／男性
希望単価：45万円
スキル：C#／C／C++
ーーーーーーーーーーーーー
"""


def test_split_single():
    blocks = split_into_blocks(SAMPLE_1)
    assert len(blocks) == 1, f"Expected 1 block, got {len(blocks)}"
    assert "H.S" in blocks[0]
    print("test_split_single: OK")


def test_split_multiple():
    blocks = split_into_blocks(SAMPLE_2)
    assert len(blocks) == 3, f"Expected 3 blocks, got {len(blocks)}: {blocks}"
    names = ["OA", "R.H", "U.H"]
    for i, name in enumerate(names):
        assert name in blocks[i], f"Expected '{name}' in block {i}"
    print("test_split_multiple: OK")


if __name__ == "__main__":
    test_split_single()
    test_split_multiple()
    print("All tests passed.")
'''

# ============================================================
# ファイル書き出し
# ============================================================
for rel_path, content in FILES.items():
    full_path = os.path.join(BASE, rel_path)
    os.makedirs(os.path.dirname(full_path), exist_ok=True)
    with open(full_path, "w", encoding="utf-8") as f:
        f.write(content)
    print(f"written: {rel_path}")

print("All files generated.")
