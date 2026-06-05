"""
skill_reader.py - 完全版
スキルシート（PDF/Word/画像）読み取り → スキル抽出 → 案件照合 → 意向確認文生成
Task14: Notionエンジニアスキル欄更新
Task15: base64/メール連携入力対応
Task16: LINE連携入力対応（base64と同じ経路）
Task17: 意向確認メール文面への自動埋め込み
"""
import os, sys, json, base64, argparse, io
try:
    sys.stdout.reconfigure(encoding='utf-8', errors='replace')
except Exception:
    pass
import requests
import pdfplumber
from docx import Document
import anthropic
from dotenv import dotenv_values

sys.path.insert(0, os.path.join(os.path.dirname(__file__), '..'))
from common.ledger import can_spend as ledger_can_spend, record as ledger_record
from common.model_config import TEXT_MODEL, VISION_MODEL

# ── 環境変数 ──────────────────────────────────────────────────────
env_path = os.path.join(os.path.dirname(__file__), '..', 'config', '.env')
config = dotenv_values(env_path)
for k, v in config.items():
    if k not in os.environ:
        os.environ[k] = v

ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
NOTION_API_KEY    = os.environ.get('NOTION_API_KEY', '')
PROJECT_DB_ID     = os.environ.get('NOTION_PROJECT_DB_ID', '343450ff-37c0-81e4-934e-f25f90284a3c')
ENGINEER_DB_ID    = os.environ.get('NOTION_ENGINEER_DB_ID', '')

NOTION_HEADERS = {
    "Authorization": f"Bearer {NOTION_API_KEY}",
    "Content-Type": "application/json",
    "Notion-Version": "2022-06-28"
}

# エンジニアDBに登録できるスキルの標準名（DBのmulti_selectの選択肢に合わせる）
VALID_SKILLS = {
    "Java", "Python", "PHP", "JavaScript", "TypeScript", "C#",
    "Node.js", "React", "Vue.js", "AWS", "インフラ", "PostgreSQL",
    "Oracle", "MySQL", "Docker", "GCP", "Go", "Ruby", "Swift",
    "Azure", "Linux", "MongoDB", "Spring"
}

SKILL_SYSTEM_PROMPT = """あなたはSES業界のスキルシート解析AIです。
入力されたスキルシートから以下をJSON形式で抽出してください。
出力はJSONのみ。前後に説明文・マークダウン記号不要。

{
  "name": "氏名（イニシャルでも可、不明なら不明）",
  "skills": ["スキル名1", "スキル名2"],
  "experience_years": {"Java": 5, "Python": 3},
  "level": "上級SE|SE|上級PG|PG",
  "summary": "一言サマリー（50文字以内）"
}

スキル名は必ず以下の標準名に統一すること:
Java, Python, PHP, JavaScript, TypeScript, C#, C言語, COBOL,
Node.js, React, Vue.js, Spring, Laravel, Ruby, Go, Swift,
AWS, Azure, GCP, Docker, Kubernetes, Linux,
Oracle, MySQL, PostgreSQL, MongoDB,
インフラ, NW設計, セキュリティ, PMO, PM, 要件定義, 基本設計, 詳細設計"""


# ── ファイル読み取り ──────────────────────────────────────────────

def extract_text_from_pdf(data: bytes) -> str:
    try:
        with pdfplumber.open(io.BytesIO(data)) as pdf:
            text = "\n".join(p.extract_text() or "" for p in pdf.pages)
        return text.strip() if len(text.strip()) > 100 else None
    except Exception:
        return None

def pdf_to_base64_image(data: bytes) -> str:
    try:
        import pypdfium2 as pdfium
        doc = pdfium.PdfDocument(data)
        page = doc[0]
        bitmap = page.render(scale=2)
        pil_img = bitmap.to_pil()
        buf = io.BytesIO()
        pil_img.save(buf, format="PNG")
        return base64.standard_b64encode(buf.getvalue()).decode()
    except Exception:
        return None

def extract_text_from_docx(data: bytes) -> str:
    doc = Document(io.BytesIO(data))
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


# ── Claude API スキル抽出 ─────────────────────────────────────────

def extract_skills_from_text(text: str) -> dict:
    model = TEXT_MODEL
    est_in = len(text[:8000]) // 4 + 500
    est_out = 1000
    if not ledger_can_spend(est_in, est_out, model):
        raise RuntimeError(f"cost_guard: skill_reader text API stopped model={model}")
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    msg = client.messages.create(
        model=model,
        max_tokens=1000,
        system=SKILL_SYSTEM_PROMPT,
        messages=[{"role": "user", "content": f"以下のスキルシートを解析してください:\n\n{text[:8000]}"}]
    )
    usage = getattr(msg, "usage", None)
    if usage:
        ledger_record(
            getattr(usage, "input_tokens", 0),
            getattr(usage, "output_tokens", 0),
            getattr(msg, "model", None) or model,
            "skill_reader",
        )
    raw = msg.content[0].text.strip().strip("```").lstrip("json").strip()
    return json.loads(raw)

def extract_skills_from_image(b64: str, mime: str = "image/png") -> dict:
    model = VISION_MODEL
    est_in = len(b64) // 4 + 500
    est_out = 1000
    if not ledger_can_spend(est_in, est_out, model):
        raise RuntimeError(f"cost_guard: skill_reader vision API stopped model={model}")
    client = anthropic.Anthropic(api_key=ANTHROPIC_API_KEY)
    msg = client.messages.create(
        model=model,
        max_tokens=1000,
        system=SKILL_SYSTEM_PROMPT,
        messages=[{
            "role": "user",
            "content": [
                {"type": "image", "source": {"type": "base64", "media_type": mime, "data": b64}},
                {"type": "text", "text": "このスキルシートを解析してください。"}
            ]
        }]
    )
    usage = getattr(msg, "usage", None)
    if usage:
        ledger_record(
            getattr(usage, "input_tokens", 0),
            getattr(usage, "output_tokens", 0),
            getattr(msg, "model", None) or model,
            "skill_reader",
        )
    raw = msg.content[0].text.strip().strip("```").lstrip("json").strip()
    return json.loads(raw)


# ── Notion エンジニアDB 更新（Task14）───────────────────────────

def update_engineer_notion(engineer_id: str, engineer_info: dict):
    """抽出したスキルでNotionエンジニアDBのスキル欄を更新"""
    extracted = set(engineer_info.get("skills", []))
    # DBのmulti_selectに存在するスキルのみ登録
    valid = [{"name": s} for s in extracted if s in VALID_SKILLS]

    payload = {"properties": {"スキル": {"multi_select": valid}}}

    # levelからDBにメモがあれば備考に追記
    summary = engineer_info.get("summary", "")
    if summary:
        payload["properties"]["備考（LINEメモ）"] = {
            "rich_text": [{"text": {"content": f"[自動解析] {summary}"}}]
        }

    r = requests.patch(
        f"https://api.notion.com/v1/pages/{engineer_id}",
        headers=NOTION_HEADERS,
        json=payload
    )
    if r.status_code == 200:
        print(f"  Notion更新完了: {len(valid)}スキル登録（{[s['name'] for s in valid]}）")
    else:
        print(f"  Notion更新失敗: {r.status_code} {r.text[:200]}")
    return r.status_code == 200


# ── Notion 案件DB 取得 ────────────────────────────────────────────

def get_active_projects(filter_keyword: str = None) -> list:
    results = []
    payload = {
        "page_size": 100,
        "filter": {"property": "ステータス", "select": {"equals": "募集中"}}
    }
    while True:
        r = requests.post(
            f"https://api.notion.com/v1/databases/{PROJECT_DB_ID}/query",
            headers=NOTION_HEADERS, json=payload
        )
        data = r.json()
        results.extend(data.get("results", []))
        if not data.get("has_more"):
            break
        payload["start_cursor"] = data["next_cursor"]

    if filter_keyword:
        kw = filter_keyword.lower()
        results = [p for p in results if kw in _get_project_name(p).lower()]
    return results

def _get_project_name(p):
    items = p["properties"].get("案件名", {}).get("title", [])
    return items[0]["plain_text"] if items else ""


# ── スキル照合 ────────────────────────────────────────────────────

def match_skills(engineer_skills: list, projects: list, engineer_price: int = None) -> list:
    eng_set = set(engineer_skills)
    results = []

    for proj in projects:
        pp = proj["properties"]
        pname   = _get_project_name(proj)
        client  = (pp.get("クライアント", {}).get("rich_text") or [{}])[0].get("plain_text", "")
        pprice  = pp.get("単価（万円）", {}).get("number")
        required = [o["name"] for o in pp.get("必要スキル", {}).get("multi_select", [])]
        optional = [o["name"] for o in pp.get("尚可スキル", {}).get("multi_select", [])]

        req_result = {s: (s in eng_set) for s in required}
        opt_result = {s: (s in eng_set) for s in optional}
        req_all_ok = all(req_result.values()) if required else True
        opt_match  = sum(1 for v in opt_result.values() if v)
        opt_total  = len(optional)
        opt_rate   = opt_match / opt_total if opt_total > 0 else 1.0
        gross      = pprice - engineer_price if pprice and engineer_price else None

        results.append({
            "project_id":     proj["id"],
            "project_name":   pname,
            "client":         client,
            "project_price":  pprice,
            "required":       req_result,
            "optional":       opt_result,
            "required_all_ok": req_all_ok,
            "opt_rate":       opt_rate,
            "opt_match":      opt_match,
            "opt_total":      opt_total,
            "gross":          gross,
            "proposable":     req_all_ok
        })

    def sort_key(r):
        if not r["proposable"]:
            return (2, 0)
        g = r["gross"]
        if g is None:
            return (1, 0)
        if 5 <= g <= 12:
            return (0, -g)
        return (1, abs(g - 8))  # 8万から離れるほど後ろ

    results.sort(key=sort_key)
    return results


# ── 意向確認メール文面生成（Task17）──────────────────────────────

def generate_iko_mail(engineer_info: dict, match_results: list,
                      engineer_price: int, affiliation_name: str = "貴社") -> str:
    """
    提案可かつ粗利ジャスト（5〜12万）の案件TOP3について意向確認メール文面を生成。
    テンプレート集v1 テンプレート1準拠。
    """
    candidates = [r for r in match_results if r["proposable"] and
                  r["gross"] is not None and 5 <= r["gross"] <= 12][:3]

    # 粗利ジャストがなければ提案可の上位3件
    if not candidates:
        candidates = [r for r in match_results if r["proposable"]][:3]

    if not candidates:
        return "（提案可能な案件がありません）"

    eng_name = engineer_info.get("name", "ご担当者")
    mails = []

    for r in candidates:
        # 必須・尚可の○×フォーマット
        req_lines = "\n".join(f" □ {s}：" for s in r["required"]) if r["required"] else " □ （必須スキルなし）"
        opt_lines = "\n".join(f" □ {s}：" for s in r["optional"]) if r["optional"] else " □ （尚可スキルなし）"

        price_str = f"{r['project_price']}万円" if r["project_price"] else "応相談"

        mail = f"""件名: {eng_name}様 案件ご検討のお願い

{affiliation_name} ご担当者様

いつもお世話になっております。

人員のご紹介ありがとうございます。
下記案件いかがでしょうか。
ご検討いただけますと幸いです。

また、エントリーいただける場合下記2点ご教授いただけますと幸いです。
・並行状況
・必須、尚可の○×

━━━━━━━━━━━━━━━━━━
■ 案件概要
━━━━━━━━━━━━━━━━━━
案件名    : {r['project_name']}
単価      : {price_str}
クライアント: {r['client']}

━━━━━━━━━━━━━━━━━━
■ ご記入フォーマット
━━━━━━━━━━━━━━━━━━
▼必須スキル（○/×）
{req_lines}
▼尚可スキル（○/×）
{opt_lines}

▼並行状況
 例）
  ・A社: 面談調整中
  ・B社: 結果待ち

何卒よろしくお願いいたします。"""
        mails.append(mail)

    return "\n\n" + ("=" * 60 + "\n").join(mails)


# ── コンソール出力 ────────────────────────────────────────────────

def print_results(engineer_info: dict, match_results: list, engineer_price: int = None):
    print("=" * 60)
    print("【スキルシート解析結果】")
    print(f"  氏名    : {engineer_info.get('name', '不明')}")
    print(f"  スキル  : {', '.join(engineer_info.get('skills', []))}")
    print(f"  レベル  : {engineer_info.get('level', '不明')}")
    if engineer_price:
        print(f"  単価    : {engineer_price}万円")
    print(f"  概要    : {engineer_info.get('summary', '')}")
    print()
    print("【案件照合結果】（粗利5〜12万を優先）")

    for r in match_results:
        ok    = "[OK]" if r["proposable"] else "[NG]"
        g_str = f"粗利{r['gross']}万" if r["gross"] is not None else "粗利不明"
        p_str = f"{r['project_price']}万" if r["project_price"] else "単価未設定"

        print(f"\n{'━'*55}")
        print(f"{ok} {r['project_name']} ({r['client']}) | {p_str} | {g_str}")

        req_str = "  ".join(f"{s}:{'○' if v else '×'}" for s, v in r["required"].items()) if r["required"] else "なし"
        print(f"  必須: {req_str}")

        if r["optional"]:
            opt_str = "  ".join(f"{s}:{'○' if v else '×'}" for s, v in r["optional"].items())
            print(f"  尚可: {opt_str} ({r['opt_match']}/{r['opt_total']})")

        if not r["proposable"]:
            ng = [s for s, v in r["required"].items() if not v]
            print(f"  → 提案不可（必須NG: {', '.join(ng)}）")
        elif r["gross"] and 5 <= r["gross"] <= 12:
            print(f"  → [S] 粗利ジャスト 提案推奨")
        elif r["gross"] and r["gross"] > 12:
            print(f"  → チャレンジ提案（単価調整要）")
        else:
            print(f"  → 提案可")

    print(f"\n{'='*60}")
    ok_count   = sum(1 for r in match_results if r["proposable"])
    just_count = sum(1 for r in match_results if r["proposable"] and r["gross"] and 5 <= r["gross"] <= 12)
    print(f"提案可: {ok_count}件  うち粗利ジャスト[S]: {just_count}件")


# ── メイン ────────────────────────────────────────────────────────

def run(file_path=None, b64_data=None, mime=None,
        engineer_price=None, filter_keyword=None,
        engineer_id=None, affiliation=None, output_mail=False):

    # 1. ファイル読み取り → base64化
    if file_path:
        with open(file_path, "rb") as f:
            raw = f.read()
        ext = os.path.splitext(file_path)[1].lower()
        mime_map = {".pdf": "application/pdf",
                    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ".doc":  "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    ".png":  "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg"}
        mime = mime_map.get(ext, "application/octet-stream")
        b64_data = base64.standard_b64encode(raw).decode()

    if not b64_data:
        print("ERROR: --file または --base64 が必要です")
        sys.exit(1)

    raw_bytes = base64.b64decode(b64_data)

    # 2. スキル抽出
    print("スキルシート解析中...")
    info = None

    if mime == "application/pdf":
        text = extract_text_from_pdf(raw_bytes)
        if text:
            info = extract_skills_from_text(text)
        else:
            print("  (テキスト薄い → 画像変換して解析)")
            b64img = pdf_to_base64_image(raw_bytes)
            if b64img:
                info = extract_skills_from_image(b64img, "image/png")
    elif "word" in (mime or ""):
        text = extract_text_from_docx(raw_bytes)
        info = extract_skills_from_text(text)
    elif mime and mime.startswith("image/"):
        info = extract_skills_from_image(b64_data, mime)
    else:
        text = raw_bytes.decode("utf-8", errors="ignore")
        info = extract_skills_from_text(text)

    if not info:
        print("ERROR: スキル抽出失敗")
        sys.exit(1)

    # 3. Notion エンジニアDB 更新（Task14）
    if engineer_id:
        print(f"Notionエンジニアスキル更新中 ({engineer_id})...")
        update_engineer_notion(engineer_id, info)

    # 4. 案件取得
    print("案件DB取得中...")
    projects = get_active_projects(filter_keyword)
    print(f"  募集中: {len(projects)}件")

    # 5. 照合
    results = match_skills(info.get("skills", []), projects, engineer_price)

    # 6. コンソール出力
    print_results(info, results, engineer_price)

    # 7. 意向確認メール文面生成（Task17）
    mail_text = generate_iko_mail(info, results, engineer_price, affiliation or "貴社")
    if output_mail:
        print("\n" + "=" * 60)
        print("【意向確認メール文面（粗利ジャスト案件TOP3）】")
        print(mail_text)

    # 8. JSON保存
    output = {"engineer": info, "engineer_price": engineer_price,
              "match_results": results, "iko_mail": mail_text}
    out_path = os.path.join(os.path.dirname(__file__), "last_result.json")
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(output, f, ensure_ascii=False, indent=2)
    print(f"\nJSON保存: {out_path}")

    return output


if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="スキルシート読み取り＆案件照合")
    parser.add_argument("--file",         help="スキルシートのファイルパス")
    parser.add_argument("--base64",       dest="b64", help="base64エンコード済みデータ")
    parser.add_argument("--mime",         help="MIMEタイプ（base64使用時）")
    parser.add_argument("--price",        type=int, help="エンジニアの希望単価（万円）")
    parser.add_argument("--filter",       help="照合案件の絞り込みキーワード")
    parser.add_argument("--engineer-id",  help="NotionエンジニアページID（スキル自動更新）")
    parser.add_argument("--affiliation",  help="所属会社名（意向確認メール宛先名）")
    parser.add_argument("--mail",         action="store_true", help="意向確認メール文面も出力")
    args = parser.parse_args()

    run(file_path=args.file, b64_data=args.b64, mime=args.mime,
        engineer_price=args.price, filter_keyword=args.filter,
        engineer_id=args.engineer_id, affiliation=args.affiliation,
        output_mail=args.mail)
