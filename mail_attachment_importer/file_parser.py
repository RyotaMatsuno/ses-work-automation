"""
file_parser.py - Excel/PDF/Word → テキスト変換モジュール
"""
import logging
from io import BytesIO

logger = logging.getLogger(__name__)


def parse_excel(data: bytes) -> str:
    """openpyxlで全シート・全セルをテキスト化"""
    try:
        import openpyxl
        wb = openpyxl.load_workbook(BytesIO(data), data_only=True)
        lines = []
        for sheet in wb.worksheets:
            lines.append(f"=== シート: {sheet.title} ===")
            for row in sheet.iter_rows():
                row_texts = []
                for cell in row:
                    if cell.value is not None:
                        row_texts.append(str(cell.value).strip())
                if row_texts:
                    lines.append("\t".join(row_texts))
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"Excel解析失敗: {e}")
        raise


def parse_pdf(data: bytes) -> str:
    """pdfplumberで全ページテキスト抽出"""
    try:
        import pdfplumber
        lines = []
        with pdfplumber.open(BytesIO(data)) as pdf:
            for i, page in enumerate(pdf.pages):
                text = page.extract_text()
                if text:
                    lines.append(f"=== ページ {i+1} ===")
                    lines.append(text)
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"PDF解析失敗: {e}")
        raise


def parse_word(data: bytes) -> str:
    """python-docxで本文テキスト化"""
    try:
        from docx import Document
        doc = Document(BytesIO(data))
        lines = [para.text for para in doc.paragraphs if para.text.strip()]
        # テーブルも取得
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    lines.append("\t".join(cells))
        return "\n".join(lines)
    except Exception as e:
        logger.error(f"Word解析失敗: {e}")
        raise


def parse_file(filename: str, ext: str, data: bytes) -> str:
    """
    拡張子に応じてパーサーを選択してテキストを返す。
    失敗時はNoneを返す。
    """
    try:
        if ext in (".xlsx", ".xls"):
            return parse_excel(data)
        elif ext == ".pdf":
            return parse_pdf(data)
        elif ext in (".docx", ".doc"):
            return parse_word(data)
        else:
            logger.warning(f"対応外形式: {ext} ({filename})")
            return None
    except Exception as e:
        logger.error(f"ファイル解析スキップ: {filename} - {e}")
        return None


if __name__ == "__main__":
    import sys
    logging.basicConfig(level=logging.INFO)
    if len(sys.argv) < 2:
        print("使い方: python file_parser.py <ファイルパス>")
        sys.exit(1)
    path = sys.argv[1]
    ext = path.rsplit(".", 1)[-1].lower()
    ext = f".{ext}"
    with open(path, "rb") as f:
        data = f.read()
    text = parse_file(path, ext, data)
    print(text[:2000] if text else "解析失敗")
